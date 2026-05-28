from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\JonatanIsaacUmulNima\source\repos\TallerMecanico\TallerMecanico\Documentacion_TallerMecanico.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, header_fill="E8EEF5"):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    style_table(table)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(body)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def main():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Documentacion Tecnica\nSistema Control de Taller Mecanico")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Aplicacion C++/CLI Windows Forms + MySQL")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_note(
        doc,
        "Proposito",
        "Este documento explica la organizacion del codigo, la responsabilidad de cada archivo y el flujo principal del sistema de taller: clientes, vehiculos, recepcion, trabajos, facturacion e historial.",
    )

    doc.add_heading("1. Resumen Del Sistema", level=1)
    doc.add_paragraph(
        "La aplicacion permite administrar un taller mecanico desde la recepcion del cliente y su vehiculo hasta el control de trabajos, uso de repuestos, cambio de estados, historial y facturacion. "
        "El sistema fue construido como aplicacion de escritorio en Visual Studio 2022 usando C++/CLI con Windows Forms y una base de datos MySQL."
    )
    add_bullet(doc, "Gestiona clientes, vehiculos, mecanicos, repuestos y tipos de servicio.")
    add_bullet(doc, "Crea ordenes de servicio para mantenimiento o reparacion.")
    add_bullet(doc, "Registra detalle del trabajo, repuestos usados, mano de obra y observaciones.")
    add_bullet(doc, "Controla estados como RECIBIDO, EN_REVISION, EN_REPARACION, LISTO, ENTREGADO y CANCELADO.")
    add_bullet(doc, "Genera facturas sobre ordenes listas o entregadas.")
    add_bullet(doc, "Envia notificacion por WhatsApp cuando una orden cambia a LISTO.")

    doc.add_heading("2. Estructura Del Proyecto", level=1)
    doc.add_paragraph(
        "El proyecto se reorganizo tomando como referencia una arquitectura sencilla por capas. En Visual Studio se muestran filtros logicos para separar formularios, conexion, clases, controladores y scripts."
    )
    add_table(
        doc,
        ["Filtro", "Contenido", "Responsabilidad"],
        [
            ["Formularios", "MenuPrincipal.h, FormClientes.h, FormRecepcion.h, FormFacturacion.h, etc.", "Pantallas de Windows Forms y eventos de botones."],
            ["Conexion", "ConfiguracionApp.h, ConexionBD.h, ConexionBD.cpp, Conexion.h", "Parametros, cadena de conexion y ejecucion centralizada de consultas."],
            ["Clases", "ClasesModelo.h", "Modelos de datos: Cliente, Vehiculo, ServicioTaller, Factura, etc."],
            ["Controladores", "ControladorTaller.h, ControladorTaller.cpp", "Capa intermedia para consultas y operaciones reutilizables."],
            ["Scripts", "actualizar_bd_taller.sql", "Cambios de base de datos, incluyendo la tabla facturas."],
        ],
        [1.15, 2.3, 3.05],
    )

    doc.add_heading("3. Capas Principales", level=1)
    doc.add_heading("3.1 ConfiguracionApp.h", level=2)
    doc.add_paragraph(
        "Centraliza datos que antes estaban dispersos, como servidor, usuario, clave, base de datos, URL de WhatsApp y API key. Esto facilita mantenimiento porque cualquier cambio de servidor o endpoint se hace en un solo archivo."
    )
    add_code(doc, 'ConfiguracionApp::CadenaConexion();')
    add_code(doc, 'ConfiguracionApp::WhatsAppUrl();')

    doc.add_heading("3.2 ConexionBD.h / ConexionBD.cpp", level=2)
    doc.add_paragraph(
        "Maneja la conexion a MySQL y ofrece metodos reutilizables para consultas, comandos INSERT/UPDATE/DELETE y valores escalares. Es equivalente al archivo ConexionBD del proyecto de restaurante, adaptado al dominio del taller."
    )
    add_bullet(doc, "AbrirConexion(): abre y reutiliza una conexion activa.")
    add_bullet(doc, "CerrarConexion(): cierra la conexion si esta abierta.")
    add_bullet(doc, "EjecutarConsultaTabla(): devuelve un DataTable para cargar DataGridView o ComboBox.")
    add_bullet(doc, "EjecutarIUD(): ejecuta INSERT, UPDATE o DELETE.")
    add_bullet(doc, "EjecutarEscalar(): devuelve un solo valor, por ejemplo LAST_INSERT_ID().")

    doc.add_heading("3.3 Conexion.h", level=2)
    doc.add_paragraph(
        "Se mantiene como puente de compatibilidad. Los formularios existentes siguen llamando a Conexion::ObtenerConexion(), pero internamente ahora usa ConexionBD::CrearConexion(). Esto evita reescribir todos los forms de una sola vez."
    )

    doc.add_heading("3.4 ClasesModelo.h", level=2)
    doc.add_paragraph(
        "Contiene clases simples que representan las entidades principales del sistema. Sirven para mover datos de forma ordenada entre formularios y controladores."
    )
    add_table(
        doc,
        ["Clase", "Representa"],
        [
            ["Cliente", "Datos personales del cliente."],
            ["Vehiculo", "Vehiculo asociado a un cliente."],
            ["Mecanico", "Empleado o tecnico asignado al servicio."],
            ["Repuesto", "Inventario usado en trabajos."],
            ["TipoServicio", "Catalogo de servicios y precios base."],
            ["ServicioTaller", "Orden de recepcion y control del trabajo."],
            ["DetalleTrabajo", "Trabajos realizados, repuestos, mano de obra y observaciones."],
            ["Factura", "Documento de cobro asociado a una orden."],
        ],
        [1.6, 4.9],
    )

    doc.add_heading("3.5 ControladorTaller.h / ControladorTaller.cpp", level=2)
    doc.add_paragraph(
        "Es la capa que concentra consultas reutilizables. La idea es que, con el tiempo, los formularios llamen a este controlador en vez de construir SQL directamente dentro de cada pantalla."
    )
    add_bullet(doc, "ListarClientes(), BuscarClientes(), InsertarCliente(), ActualizarCliente(), EliminarCliente().")
    add_bullet(doc, "ListarVehiculos(), ListarMecanicos(), ListarRepuestos(), ListarTiposServicio().")
    add_bullet(doc, "ListarServicios() y ListarFacturas().")
    add_bullet(doc, "CrearTablaFacturasSiNoExiste(), usado por el modulo de facturacion para evitar errores si la tabla no existe.")

    doc.add_heading("4. Formularios Del Sistema", level=1)
    add_table(
        doc,
        ["Formulario", "Funcion principal"],
        [
            ["MenuPrincipal.h", "Pantalla inicial. Abre cada modulo del sistema."],
            ["FormClientes.h", "CRUD de clientes: guardar, modificar, eliminar, buscar y listar."],
            ["FormVehiculos.h", "CRUD de vehiculos relacionados con clientes."],
            ["FormMecanicos.h", "CRUD de mecanicos y especialidades."],
            ["FormRepuestos.h", "Inventario de repuestos, precios y stock."],
            ["FormTipoServicio.h", "Catalogo de servicios y precios base."],
            ["FormRecepcion.h", "Crea ordenes, actualiza estados y registra detalle de trabajos."],
            ["FormHistorial.h", "Consulta historica de ordenes y sus detalles."],
            ["FormFacturacion.h", "Genera facturas de ordenes LISTO o ENTREGADO."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("5. Flujo De Recepcion Y Trabajo", level=1)
    doc.add_paragraph("El flujo principal del taller inicia en el modulo Recepcion / Orden:")
    add_number(doc, "Seleccionar cliente.")
    add_number(doc, "Seleccionar vehiculo del cliente.")
    add_number(doc, "Seleccionar tipo de servicio y mecanico asignado.")
    add_number(doc, "Ingresar fecha, kilometraje, estado y problema reportado.")
    add_number(doc, "Crear la orden en la tabla servicios.")
    add_number(doc, "Agregar detalle de trabajo, repuesto, cantidad, mano de obra y observacion.")
    add_number(doc, "Recalcular el total de la orden.")
    add_number(doc, "Actualizar estado hasta LISTO o ENTREGADO.")

    add_note(
        doc,
        "Control de estado",
        "Cuando la orden cambia a LISTO, el sistema obtiene el telefono del cliente, prepara el numero de WhatsApp y envia una notificacion mediante el endpoint configurado.",
    )

    doc.add_heading("6. Facturacion", level=1)
    doc.add_paragraph(
        "El modulo de facturacion trabaja con ordenes que ya estan en estado LISTO o ENTREGADO. Esto evita facturar servicios que aun estan en revision o reparacion."
    )
    add_bullet(doc, "Carga la orden seleccionada con cliente, vehiculo, servicio, mecanico y subtotal.")
    add_bullet(doc, "Permite aplicar descuento.")
    add_bullet(doc, "Permite escoger metodo de pago: Efectivo, Tarjeta, Transferencia o Credito.")
    add_bullet(doc, "Guarda una factura por orden usando una llave unica sobre id_servicio.")
    add_bullet(doc, "Lista facturas emitidas y permite consultar el detalle de trabajos asociados.")

    doc.add_heading("7. Base De Datos", level=1)
    doc.add_paragraph(
        "La base de datos principal se llama taller_mecanico. Las tablas base son clientes, vehiculos, mecanico, tiposervicio, repuestos, servicios y detalletrabajo. Se agrego facturas para el nuevo modulo de cobro."
    )
    add_table(
        doc,
        ["Tabla", "Descripcion"],
        [
            ["clientes", "Informacion del cliente: NIT, nombre, direccion, correo y telefono."],
            ["vehiculos", "Vehiculos vinculados a un cliente."],
            ["mecanico", "Mecanicos disponibles y especialidades."],
            ["tiposervicio", "Servicios de mantenimiento o reparacion con precio base."],
            ["repuestos", "Inventario, marca, precio y stock."],
            ["servicios", "Orden principal de recepcion y seguimiento."],
            ["detalletrabajo", "Trabajos realizados y repuestos usados por servicio."],
            ["facturas", "Facturas emitidas sobre ordenes de servicio."],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("8. WhatsApp", level=1)
    doc.add_paragraph(
        "La notificacion por WhatsApp esta integrada en FormRecepcion.h. Se ejecuta al actualizar una orden a LISTO, siempre que el estado anterior no fuera LISTO. Esto evita enviar el mismo mensaje repetidamente si se vuelve a guardar el mismo estado."
    )
    add_bullet(doc, "Obtiene cliente, telefono, placa y total desde la base de datos.")
    add_bullet(doc, "Limpia el numero y le agrega el prefijo 502 si tiene ocho digitos.")
    add_bullet(doc, "Agrega @s.whatsapp.net al final del identificador.")
    add_bullet(doc, "Envia JSON con number y text usando la API configurada.")
    add_code(doc, '{"number":"502XXXXXXXX@s.whatsapp.net","text":"Su vehiculo ya esta LISTO..."}')

    doc.add_heading("9. Recomendaciones De Mantenimiento", level=1)
    add_bullet(doc, "Mover gradualmente las consultas SQL restantes desde los formularios hacia ControladorTaller.")
    add_bullet(doc, "Cambiar la clave de MySQL y la API key si el proyecto sera entregado o compartido.")
    add_bullet(doc, "Validar stock antes de descontar repuestos para evitar cantidades negativas.")
    add_bullet(doc, "Agregar usuarios y permisos si el sistema sera usado por varias personas.")
    add_bullet(doc, "Crear respaldos periodicos de la base de datos taller_mecanico.")
    add_bullet(doc, "Agregar reportes imprimibles de factura y orden de servicio si se requiere entrega fisica.")

    doc.add_heading("10. Como Ejecutar El Proyecto", level=1)
    add_number(doc, "Abrir la solucion TallerMecanico.sln en Visual Studio 2022.")
    add_number(doc, "Verificar que la referencia MySql.Data este activa.")
    add_number(doc, "Ejecutar actualizar_bd_taller.sql en MySQL si aun no se han creado los campos nuevos o la tabla facturas.")
    add_number(doc, "Compilar la solucion en Debug x64.")
    add_number(doc, "Ejecutar con F5 y probar primero Clientes, Vehiculos, Tipos de servicio, Repuestos y Mecanicos.")
    add_number(doc, "Luego probar Recepcion, Historial y Facturacion.")

    add_note(
        doc,
        "Estado actual",
        "El proyecto compila correctamente despues de la reorganizacion. La estructura nueva mejora la legibilidad sin eliminar la compatibilidad con los formularios actuales.",
    )

    doc.add_heading("11. Resumen Para Exponer", level=1)
    doc.add_paragraph(
        "El sistema se divide en formularios para la interfaz, una capa de conexion para MySQL, clases modelo para representar datos y un controlador que centraliza operaciones reutilizables. "
        "La recepcion crea ordenes de servicio, el detalle controla lo realizado en el vehiculo, historial permite consultar trabajos anteriores y facturacion genera el cobro final. "
        "La integracion con WhatsApp automatiza la comunicacion con el cliente cuando el vehiculo esta listo."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Sistema Control de Taller Mecanico")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()

